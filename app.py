import streamlit as st
import pandas as pd
import os
import io
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="Gestión Minera Avanzada", layout="wide")

# --- 1. GESTIÓN DE ESTADO ---
if 'bases_datos' not in st.session_state:
    st.session_state['bases_datos'] = {}

def detectar_columnas_meses(df):
    prefijos = [("jan", "ene"), ("feb", "feb"), ("mar", "mar"), ("apr", "abr"), 
                ("may", "may"), ("jun", "jun"), ("jul", "jul"), ("aug", "ago"), 
                ("sep", "sep"), ("oct", "oct"), ("nov", "nov"), ("dec", "dic")]
    columnas = []
    for pref_en, pref_es in prefijos:
        match = next((c for c in df.columns if str(c).strip().lower().startswith(pref_en) or str(c).strip().lower().startswith(pref_es)), None)
        if match: columnas.append(match)
    return columnas if len(columnas) == 12 else []

# --- 2. MOTORES MATEMÁTICOS ---
def aplicar_fit_forecast(row, alpha, delta, x_meses, col_budget_fy, col_meses):
    try: budget_fy = float(row[col_budget_fy]) if pd.notna(row[col_budget_fy]) else 0.0
    except: budget_fy = 0.0

    if budget_fy <= 0.01:
        return pd.Series({f"Factor_FIT_{col_meses[i]}": 1.0 for i in range(x_meses, 12)})

    presupuesto_medio_mensual = budget_fy / 12.0
    F, T, FIT = 0, 0, 0

    for t in range(x_meses):
        try: actual = float(row[col_meses[t]]) if pd.notna(row[col_meses[t]]) else 0
        except: actual = 0
        
        eficiencia_cruda = actual / presupuesto_medio_mensual
        eficiencia = min(max(eficiencia_cruda, 0.1), 3.0) 

        if t == 0: 
            F, T = eficiencia, 0; FIT = F + T
        else:
            F_nuevo = FIT + alpha * (eficiencia - FIT)
            F_nuevo = min(max(F_nuevo, 0.3), 2.5) 
            T_nuevo = T + delta * (F_nuevo - FIT)
            T_nuevo = min(max(T_nuevo, -0.05), 0.05) 
            F, T = F_nuevo, T_nuevo
            FIT = F + T

    factores_futuros = {}
    for step, i in enumerate(range(x_meses, 12)):
        factor_proyectado = F + (step + 1) * T
        factor_proyectado = min(max(factor_proyectado, 0.4), 2.0) 
        factores_futuros[f"Factor_FIT_{col_meses[i]}"] = factor_proyectado
        
    return pd.Series(factores_futuros)

def aplicar_fit_quinquenal(serie_historica, alpha, delta, meses_a_proyectar=60):
    if len(serie_historica) == 0: return [0] * meses_a_proyectar
    F, T = serie_historica[0], 0
    FIT = F + T
    for t in range(1, len(serie_historica)):
        F_nuevo = FIT + alpha * (serie_historica[t] - FIT)
        T_nuevo = T + delta * (F_nuevo - FIT)
        F, T = F_nuevo, T_nuevo
        FIT = F + T
    return [max(0, F + step * T) for step in range(1, meses_a_proyectar + 1)]

# --- 3. MOTOR DE EXPORTACIÓN (WORD) MEJORADO ---
def generar_reporte_word(titulo, kpis, parametros, conclusiones, insight_dinamico=None):
    doc = Document()
    titulo_doc = doc.add_heading(f'Informe Técnico: {titulo}', 0)
    titulo_doc.alignment = 1 
    
    doc.add_paragraph('Generado por: Sistema Predictivo de Gestión Minera Avanzada').alignment = 1
    doc.add_paragraph('_' * 70).alignment = 1
    
    # Tabla elegante para KPIs
    doc.add_heading('1. Resumen Ejecutivo (Métricas Clave)', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid' 
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Indicador Financiero'
    hdr_cells[1].text = 'Valor Proyectado'
    for key, value in kpis.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)
        
    doc.add_heading('2. Parametrización del Modelo', level=1)
    for key, value in parametros.items():
        doc.add_paragraph(f'{key}: {value}', style='List Bullet')
        
    doc.add_heading('3. Metodología Aplicada', level=1)
    doc.add_paragraph(conclusiones)
    
    # Módulo exclusivo de inteligencia de negocios
    if insight_dinamico:
        doc.add_heading('4. Insight Estratégico (Hallazgo Automático)', level=1)
        p = doc.add_paragraph()
        p.add_run(insight_dinamico).bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. MENÚ ---
st.sidebar.title("Plataforma Directiva")
menu = st.sidebar.radio("Módulos del Sistema:", ("Gestión de Datos", "Forecast", "Budget Quinquenal"))

if menu == "Gestión de Datos":
    st.title("📂 Repositorio Central")
    archivo_subido = st.file_uploader("Subir base de datos (.xlsx)", type=["xlsx", "xls"])
    if archivo_subido:
        try:
            dict_hojas = pd.read_excel(archivo_subido, sheet_name=None)
            for n, df in dict_hojas.items():
                if not df.empty and any("Unnamed:" in str(col) for col in df.columns):
                    df.columns = df.iloc[0].astype(str).str.strip()
                    df = df.iloc[1:].reset_index(drop=True)
                dict_hojas[n] = df
            st.session_state['bases_datos'][archivo_subido.name] = dict_hojas
            st.success("Archivo procesado.")
        except Exception as e: st.error(f"Error: {e}")

    if st.session_state['bases_datos']:
        for nombre, hojas in list(st.session_state['bases_datos'].items()):
            c1, c2 = st.columns([4, 1])
            with c1: st.info(f"**{nombre}** | Pestañas: {', '.join(list(hojas.keys()))}")
            with c2:
                if st.button("🗑️ Eliminar", key=f"del_{nombre}"): del st.session_state['bases_datos'][nombre]; st.rerun()

elif menu == "Forecast":
    st.title("📈 Forecast Dinámico")
    if not st.session_state['bases_datos']: st.warning("Carga una base de datos.")
    else:
        archivo_sel = st.selectbox("Archivo Base:", list(st.session_state['bases_datos'].keys()))
        hoja_target = st.selectbox("Pestaña a Proyectar:", list(st.session_state['bases_datos'][archivo_sel].keys()))
        df_base = st.session_state['bases_datos'][archivo_sel][hoja_target].copy()
        
        c1, c2, c3 = st.columns(3)
        with c1: X_meses = st.slider("Meses REALES:", 1, 11, 5)
        with c2: alpha = st.slider("Sensibilidad Pronóstico (α):", 0.0, 1.0, 0.5, 0.05)
        with c3: delta = st.slider("Sensibilidad Tendencia (δ):", 0.0, 1.0, 0.3, 0.05)
            
        col_budget_fy = next((c for c in df_base.columns if "BUDGET FY" in str(c).upper() or ("BUDGET" in str(c).upper() and "BYTD" not in str(c).upper())), None)
        columnas_meses = detectar_columnas_meses(df_base)
        
        if col_budget_fy and len(columnas_meses) == 12:
            df_factores = df_base.apply(aplicar_fit_forecast, axis=1, args=(alpha, delta, X_meses, col_budget_fy, columnas_meses))
            df_base = pd.concat([df_base, df_factores], axis=1)
            columnas_panorama = []
            for idx, col_mes in enumerate(columnas_meses):
                nombre_final = f"{col_mes} (Final)"
                try: val = pd.to_numeric(df_base[col_mes], errors='coerce').fillna(0)
                except: val = 0
                df_base[nombre_final] = val if idx < X_meses else val * df_base[f"Factor_FIT_{col_mes}"]
                columnas_panorama.append(nombre_final)
                
            total_planificado = pd.to_numeric(df_base[col_budget_fy], errors='coerce').fillna(0).sum()
            total_estimado = sum([df_base[c].sum() for c in columnas_panorama])
            variacion = total_estimado - total_planificado
            porcentaje_var = (variacion/total_planificado)*100 if total_planificado>0 else 0
            
            st.markdown("### Dashboard Ejecutivo")
            k1, k2, k3 = st.columns(3)
            k1.metric("Budget FY", f"USD {total_planificado:,.0f}")
            k2.metric("Forecast FIT", f"USD {total_estimado:,.0f}")
            k3.metric("Varianza", f"USD {variacion:,.0f}", f"{porcentaje_var:.2f}%", delta_color="inverse")
            
            df_grafico = pd.DataFrame({'Original': [pd.to_numeric(df_base[m], errors='coerce').fillna(0).sum() for m in columnas_meses], 'Proyectado': [df_base[c].sum() for c in columnas_panorama]}, index=[f"{i+1:02d}. {str(m).split('-')[0].strip()}" for i, m in enumerate(columnas_meses)])
            st.line_chart(df_grafico, use_container_width=True) # Gráfico cambiado a Líneas

            st.markdown("### Matriz de Estimaciones")
            cols_id = [c for c in df_base.columns if c not in columnas_meses and "Factor" not in c and "(Final)" not in c]
            st.dataframe(df_base[cols_id[:4] + columnas_panorama], use_container_width=True)

            col_d1, col_d2 = st.columns(2)
            df_export = df_base[cols_id[:4] + columnas_panorama].copy()
            buffer_excel = io.BytesIO()
            with pd.ExcelWriter(buffer_excel, engine='openpyxl') as writer: df_export.to_excel(writer, index=False)
            buffer_excel.seek(0)
            col_d1.download_button("📊 Descargar Excel", buffer_excel, f"Forecast_M{X_meses}.xlsx")

            kpis_doc = {"Presupuesto Original": f"USD {total_planificado:,.2f}", "Forecast": f"USD {total_estimado:,.2f}", "Desviación": f"USD {variacion:,.2f} ({porcentaje_var:.2f}%)"}
            buffer_w = generar_reporte_word("Forecast Dinámico", kpis_doc, {"Meses Reales": X_meses, "Alpha": alpha, "Delta": delta}, "Cálculo mediante FIT con límites de tolerancia operativa.")
            col_d2.download_button("📄 Descargar Reporte Word", buffer_w, "Informe_Forecast.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else: st.error("Estructura inválida.")

elif menu == "Budget Quinquenal":
    st.title("📅 Budget Quinquenal & Sensibilidades")
    if not st.session_state['bases_datos']: st.warning("Carga información histórica.")
    else:
        archivo_sel = st.selectbox("Archivo Base:", list(st.session_state['bases_datos'].keys()))
        hojas_hist = st.multiselect("Pestañas de Entrenamiento:", list(st.session_state['bases_datos'][archivo_sel].keys()))
        
        c1, c2 = st.columns(2)
        with c1: alpha_q = st.slider("Sensibilidad (α):", 0.0, 1.0, 0.4, 0.05)
        with c2: delta_q = st.slider("Aceleración (δ):", 0.0, 1.0, 0.2, 0.05)
        
        if st.button("🚀 Ejecutar Proyección Quinquenal"):
            if not hojas_hist: st.error("Selecciona pestañas.")
            else:
                with st.spinner("Procesando..."):
                    df_base = st.session_state['bases_datos'][archivo_sel][hojas_hist[0]].copy()
                    cols_id = [c for c in df_base.columns if c not in detectar_columnas_meses(df_base)]
                    df_resultado = df_base[cols_id[:5]].copy()
                    
                    meses_base = ["Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
                    años_futuros = [2027, 2028, 2029, 2030, 2031]
                    matriz_proyecciones = []
                    
                    for idx, row in df_resultado.iterrows():
                        serie_hist = []
                        for hoja in hojas_hist:
                            df_hist = st.session_state['bases_datos'][archivo_sel][hoja]
                            meses_hist = detectar_columnas_meses(df_hist)
                            if len(meses_hist) == 12:
                                for m in meses_hist:
                                    try: val = float(df_hist.iloc[idx][m]) if pd.notna(df_hist.iloc[idx][m]) else 0
                                    except: val = 0
                                    serie_hist.append(val)
                        matriz_proyecciones.append(aplicar_fit_quinquenal(serie_hist, alpha_q, delta_q, 60))
                    
                    df_matriz = pd.DataFrame(matriz_proyecciones)
                    col_idx = 0; cols_fy = []; cols_2027 = []
                    for año in años_futuros:
                        cols_año = []
                        for mes in meses_base:
                            nombre_col = f"{mes} {año}"
                            df_resultado[nombre_col] = df_matriz[col_idx]
                            cols_año.append(nombre_col)
                            if año == 2027: cols_2027.append(nombre_col)
                            col_idx += 1
                        fy_name = f"FY {año}"
                        df_resultado[fy_name] = df_resultado[cols_año].sum(axis=1)
                        cols_fy.append(fy_name)
                    
                    st.session_state['df_q'] = df_resultado
                    st.session_state['c_27'] = cols_2027
                    st.session_state['c_fy'] = cols_fy
                    st.session_state['c_id'] = cols_id[:5]
        
        if 'df_q' in st.session_state:
            df_q, c_27, c_fy, c_id = st.session_state['df_q'], st.session_state['c_27'], st.session_state['c_fy'], st.session_state['c_id']
            
            st.markdown("---")
            st.markdown("### ⚙️ Módulo de Sensibilidades y Ponderación")
            s1, s2, s3 = st.columns(3)
            peso_comb = s1.number_input("% Peso Combustible", 0, 100, 20)
            peso_div = s2.number_input("% Peso Divisas", 0, 100, 35)
            peso_mo = s3.number_input("% Peso Mano de Obra", 0, 100, 30)
            
            st.markdown("**Simulador de Shocks Macroeconómicos:**")
            col_s1, col_s2, col_s3 = st.columns(3)
            var_comb = col_s1.slider("🛢️ Δ Combustible (%)", -50, 50, 0)
            var_div = col_s2.slider("💱 Δ Divisas (%)", -50, 50, 0)
            var_mo = col_s3.slider("👷 Δ Mano de Obra (%)", -50, 50, 0)
            
            # Matemáticas de impacto
            factor_impacto = 1 + ((peso_comb/100) * (var_comb/100)) + ((peso_div/100) * (var_div/100)) + ((peso_mo/100) * (var_mo/100))
            df_sensibilizado = df_q.copy()
            for col in c_27 + c_fy: df_sensibilizado[col] = df_sensibilizado[col] * factor_impacto
                
            st.markdown("### Dashboard Quinquenal: Escenario Base vs. Sensibilizado")
            totales_base = [df_q[fy].sum() for fy in c_fy]
            totales_sens = [df_sensibilizado[fy].sum() for fy in c_fy]
            suma_base, suma_sens = sum(totales_base), sum(totales_sens)
            impacto_neto = suma_sens - suma_base
            porc_impacto = (impacto_neto/suma_base)*100 if suma_base>0 else 0
            
            k1, k2, k3 = st.columns(3)
            k1.metric("Costo Quinquenio (Base)", f"USD {suma_base:,.0f}")
            k2.metric("Costo Quinquenio (Sensibilizado)", f"USD {suma_sens:,.0f}")
            k3.metric("Impacto Neto (Brecha)", f"USD {impacto_neto:,.0f}", f"{porc_impacto:.2f}%", delta_color="inverse")
            
            # GRÁFICO DE LÍNEAS PARA TENDENCIA A 5 AÑOS
            df_graf_q = pd.DataFrame({'Escenario Base': totales_base, 'Escenario Sensibilizado': totales_sens}, index=[fy.replace("FY ", "") for fy in c_fy])
            st.line_chart(df_graf_q, use_container_width=True)

            # MATRIZ VISIBLE A PETICIÓN
            st.markdown("### Matriz de Estimaciones Quinquenales")
            st.dataframe(df_sensibilizado[c_id + c_27 + c_fy], use_container_width=True)

            st.markdown("### Exportación de Reportes")
            col_d1, col_d2 = st.columns(2)
            buffer_q = io.BytesIO()
            with pd.ExcelWriter(buffer_q, engine='openpyxl') as writer: df_sensibilizado[c_id + c_27 + c_fy].to_excel(writer, index=False)
            buffer_q.seek(0)
            col_d1.download_button("📊 Descargar Matriz Quinquenal", buffer_q, "Budget_Quinquenal.xlsx")
            
            # CÁLCULO DE INSIGHT DINÁMICO PARA WORD
            impactos = {"Combustible": abs((peso_comb/100)*(var_comb/100)), "Divisas": abs((peso_div/100)*(var_div/100)), "Mano de Obra": abs((peso_mo/100)*(var_mo/100))}
            mayor_impacto = max(impactos, key=impactos.get)
            if impactos[mayor_impacto] > 0:
                insight_txt = f"💡 ALERTA ESTRATÉGICA: La variable macroeconómica con mayor impacto en la desviación de este presupuesto es '{mayor_impacto}'. Esto se debe a su alta ponderación en la estructura de costos combinada con la magnitud del shock simulado. Se recomienda focalizar políticas de cobertura (hedging) en este ítem."
            else:
                insight_txt = "💡 El escenario quinquenal se mantiene en su tendencia natural. No se registran perturbaciones o shocks macroeconómicos activos en la simulación actual."

            kpis_q = {"Proyección Base (5 años)": f"USD {suma_base:,.2f}", "Proyección Sensibilizada": f"USD {suma_sens:,.2f}", "Impacto Neto Absoluto": f"USD {impacto_neto:,.2f} ({porc_impacto:.2f}%)"}
            params_q = {"Alpha": alpha_q, "Delta": delta_q, "Combustible": f"Peso {peso_comb}% | Var {var_comb}%", "Divisas": f"Peso {peso_div}% | Var {var_div}%", "MO": f"Peso {peso_mo}% | Var {var_mo}%"}
            buffer_word_q = generar_reporte_word("Presupuesto Quinquenal 2027-2031", kpis_q, params_q, "Generación mediante modelo FIT y ponderación macroeconómica.", insight_txt)
            col_d2.download_button("📄 Descargar Informe Técnico", buffer_word_q, "Informe_Quinquenal.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")